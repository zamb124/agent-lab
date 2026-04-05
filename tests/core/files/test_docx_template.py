"""Тесты сервиса шаблонизации DOCX (core/files/docx_template)."""

from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from core.files.docx_template import (
    DocxTemplateContextError,
    DocxTemplateInvalidError,
    DocxTemplater,
    DocxTemplateSyntaxError,
    render_docx_template_bytes,
)
from core.files.docx_template.normalize import normalize_template_context
from core.files.models import FileRecord, FileStatus
from core.files.writer import FileWriteError


def _paragraph_template_bytes(paragraph_text: str) -> bytes:
    buf = BytesIO()
    doc = Document()
    doc.add_paragraph(paragraph_text)
    doc.save(buf)
    return buf.getvalue()


def _file_info(tmp_path: Path, paragraph_text: str, name: str = "tpl.docx") -> dict:
    p = tmp_path / name
    p.write_bytes(_paragraph_template_bytes(paragraph_text))
    return {"name": name, "path": str(p)}


def _read_first_paragraph(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    return doc.paragraphs[0].text


def _t() -> DocxTemplater:
    return DocxTemplater()


@pytest.mark.asyncio
async def test_render_simple_substitution(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "Hello {{ name }}")
    out = await _t().fill(finfo, {"name": "World"})
    assert b"PK" in out[:4]
    assert _read_first_paragraph(out) == "Hello World"


@pytest.mark.asyncio
async def test_render_nested_and_date_iso(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "{{ user.name }} {{ d }}")
    out = await _t().fill(
        finfo,
        {"user": {"name": "Ann"}, "d": date(2026, 4, 5)},
        date_iso=True,
    )
    text = _read_first_paragraph(out)
    assert "Ann" in text
    assert "2026-04-05" in text


@pytest.mark.asyncio
async def test_render_datetime_iso(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "{{ ts }}")
    dt = datetime(2026, 4, 5, 12, 30, 0)
    out = await _t().fill(finfo, {"ts": dt})
    assert "2026-04-05T12:30:00" in _read_first_paragraph(out)


@pytest.mark.asyncio
async def test_strict_missing_variable_raises(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "Hello {{ name }}")
    with pytest.raises(DocxTemplateContextError) as exc_info:
        await _t().fill(
            finfo,
            {},
            strict=True,
        )
    assert exc_info.value.payload.get("missing_variables") == ["name"]


@pytest.mark.asyncio
async def test_strict_extra_key_raises(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "Hello {{ name }}")
    with pytest.raises(DocxTemplateContextError) as exc_info:
        await _t().fill(
            finfo,
            {"name": "x", "unused": 1},
            strict=True,
        )
    assert exc_info.value.payload.get("extra_keys") == ["unused"]


@pytest.mark.asyncio
async def test_non_strict_allows_missing_and_extra(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "Hello {{ name }}")
    out = await _t().fill(
        finfo,
        {"extra_only": 1},
        strict=False,
    )
    assert _read_first_paragraph(out) == "Hello "


@pytest.mark.asyncio
async def test_invalid_zip_raises(tmp_path: Path) -> None:
    p = tmp_path / "bad.docx"
    p.write_bytes(b"not-a-zip")
    finfo = {"name": "bad.docx", "path": str(p)}
    with pytest.raises(DocxTemplateInvalidError):
        await _t().fill(finfo, {"a": 1})


@pytest.mark.asyncio
async def test_empty_template_raises(tmp_path: Path) -> None:
    p = tmp_path / "empty.docx"
    p.write_bytes(b"")
    finfo = {"name": "empty.docx", "path": str(p)}
    with pytest.raises(DocxTemplateInvalidError):
        await _t().fill(finfo, {})


@pytest.mark.asyncio
async def test_jinja_syntax_error_wraps(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "broken {{ x")
    with pytest.raises(DocxTemplateSyntaxError):
        await _t().fill(finfo, {"x": 1})


def test_normalize_rejects_callable() -> None:
    with pytest.raises(DocxTemplateContextError):
        normalize_template_context(
            {"fn": lambda: 1},
            date_iso=True,
        )


@pytest.mark.asyncio
async def test_fill_and_create_rejects_non_docx_output_name(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "{{ a }}")
    with pytest.raises(FileWriteError, match="\\.docx"):
        await _t().fill_and_create(
            file_ref=finfo,
            context={"a": "x"},
            output_original_name="out.txt",
        )


@pytest.mark.asyncio
async def test_deterministic_output_size_stable(tmp_path: Path) -> None:
    finfo = _file_info(tmp_path, "V={{ v }}")
    out1 = await _t().fill(finfo, {"v": "1"})
    out2 = await _t().fill(finfo, {"v": "1"})
    assert len(out1) == len(out2)
    assert out1 == out2


@pytest.mark.asyncio
async def test_fill_accepts_file_record(monkeypatch: pytest.MonkeyPatch) -> None:
    tpl = _paragraph_template_bytes("Hello {{ name }}")

    async def fake_read_stored(file_id: str) -> tuple[bytes, str]:
        assert file_id == "rec1"
        return tpl, "tpl.docx"

    monkeypatch.setattr(
        "core.files.reader.service._read_stored_file_by_id",
        fake_read_stored,
    )
    record = FileRecord(
        file_id="rec1",
        provider="p",
        original_name="tpl.docx",
        s3_key="k",
        s3_bucket="b",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        file_size=len(tpl),
        status=FileStatus.READY,
    )
    out = await _t().fill(record, {"name": "World"})
    assert _read_first_paragraph(out) == "Hello World"


def test_render_docx_template_bytes_low_level() -> None:
    raw = _paragraph_template_bytes("{{ x }}")
    out = render_docx_template_bytes(raw, {"x": "y"})
    assert _read_first_paragraph(out) == "y"
