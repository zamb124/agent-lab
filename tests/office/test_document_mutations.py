from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from apps.office.services.document_mutations import (
    DocumentMutationError,
    append_text_to_document,
    replace_text_in_document,
    update_spreadsheet_cells,
)


def _docx_bytes(*paragraphs: str) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _xlsx_bytes() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws["A1"] = "hello old"
    ws["B2"] = 10
    out = BytesIO()
    wb.save(out)
    return out.getvalue()


def test_replace_text_in_txt_case_insensitive():
    data, count = replace_text_in_document(
        data=b"Hello old\nhello OLD\n",
        original_name="notes.txt",
        find="old",
        replace="new",
        match_case=False,
    )

    assert data == b"Hello new\nhello new\n"
    assert count == 2


def test_replace_and_append_docx_roundtrip():
    data, count = replace_text_in_document(
        data=_docx_bytes("Project status: old"),
        original_name="report.docx",
        find="old",
        replace="ready",
        match_case=True,
    )
    appended = append_text_to_document(
        data=data,
        original_name="report.docx",
        text="Next line",
    )
    doc = Document(BytesIO(appended))

    assert count == 1
    assert [p.text for p in doc.paragraphs] == ["Project status: ready", "Next line"]


def test_replace_and_update_xlsx_roundtrip():
    data, count = replace_text_in_document(
        data=_xlsx_bytes(),
        original_name="data.xlsx",
        find="old",
        replace="new",
        match_case=True,
    )
    updated = update_spreadsheet_cells(
        data=data,
        original_name="data.xlsx",
        sheet="Data",
        cells={"B2": 20, "C3": "done"},
    )
    wb = load_workbook(BytesIO(updated))
    ws = wb["Data"]

    assert count == 1
    assert ws["A1"].value == "hello new"
    assert ws["B2"].value == 20
    assert ws["C3"].value == "done"


def test_update_csv_expands_rows_and_columns():
    data = update_spreadsheet_cells(
        data=b"a,b\n",
        original_name="data.csv",
        sheet=None,
        cells={"C2": "x", "A3": 5},
    )

    assert data.decode("utf-8").splitlines() == ["a,b", ",,x", "5"]


def test_unsupported_extension_is_rejected():
    with pytest.raises(DocumentMutationError):
        append_text_to_document(data=b"abc", original_name="archive.zip", text="x")
