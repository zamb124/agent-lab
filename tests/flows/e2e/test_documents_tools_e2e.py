"""E2E coverage for flow document tools over real flows, Office BFF and files storage."""

from __future__ import annotations

import asyncio
import base64
import uuid
from io import BytesIO
from typing import Any, cast

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

pytestmark = [
    pytest.mark.real_taskiq,
    pytest.mark.timeout(180, func_only=True),
    pytest.mark.usefixtures("frontend_service"),
]
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _headers(auth_headers_system: dict[str, str]) -> dict[str, str]:
    return {**auth_headers_system, "X-Platform-Namespace": "default"}


def _docx_bytes(*paragraphs: str) -> bytes:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend((cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx_bytes() -> bytes:
    wb = Workbook()
    ws = wb.active
    if ws is None:
        raise RuntimeError("Workbook.active is required")
    ws.title = "Sheet1"
    ws["A1"] = "OLD_CELL"
    ws["B2"] = "keep"
    out = BytesIO()
    wb.save(out)
    return out.getvalue()


async def _create_catalog(office_client_http, headers: dict[str, str], unique_id: str) -> str:
    response = await office_client_http.post(
        "/documents/api/v1/catalogs",
        headers=headers,
        json={"title": f"documents-tools-{unique_id}", "is_public": False},
    )
    assert response.status_code == 200, response.text
    return response.json()["catalog_id"]


async def _create_flow(flows_client_http, headers: dict[str, str], flow_id: str) -> None:
    response = await flows_client_http.post(
        "/flows/api/v1/flows/",
        headers=headers,
        json={
            "flow_id": flow_id,
            "name": f"Documents tools {flow_id}",
            "entry": "main",
            "nodes": {
                "main": {
                    "type": "llm_node",
                    "prompt": "Use the provided document tools to edit the attached file.",
                    "tools": [
                        "documents_open_file",
                        "documents_replace_text",
                        "documents_append_text",
                        "documents_update_cells",
                    ],
                }
            },
            "edges": [{"from_node": "main", "to_node": None}],
        },
    )
    assert response.status_code == 200, response.text


async def _send_file_message(
    flows_client_http,
    headers: dict[str, str],
    *,
    flow_id: str,
    context_id: str,
    file_name: str,
    mime_type: str,
    file_bytes: bytes,
) -> dict[str, Any]:
    response = await flows_client_http.post(
        f"/flows/api/v1/{flow_id}",
        headers=headers,
        json={
            "jsonrpc": "2.0",
            "id": f"test-{uuid.uuid4().hex}",
            "method": "message/send",
            "params": {
                "message": {
                    "messageId": str(uuid.uuid4()),
                    "role": "user",
                    "contextId": context_id,
                    "parts": [
                        {"kind": "text", "text": "Edit the attached document."},
                        {
                            "kind": "file",
                            "file": {
                                "bytes": base64.b64encode(file_bytes).decode("ascii"),
                                "name": file_name,
                                "mimeType": mime_type,
                            },
                        },
                    ],
                }
            },
        },
    )
    assert response.status_code == 200, response.text
    data = response.json()
    task = data.get("result", {})
    assert task.get("status", {}).get("state") == "completed", data
    return data


def _file_item_dict(item: object) -> dict[str, Any] | None:
    if isinstance(item, dict):
        return cast(dict[str, Any], item)
    dump = getattr(item, "model_dump", None)
    if callable(dump):
        dumped = dump(mode="json", exclude_none=True)
        if isinstance(dumped, dict):
            return cast(dict[str, Any], dumped)
    return None


def _file_document(item: dict[str, Any]) -> dict[str, Any] | None:
    capabilities = item.get("capabilities")
    document = item.get("document")
    if isinstance(document, dict):
        return document
    if isinstance(capabilities, dict):
        cap_document = capabilities.get("document")
        if isinstance(cap_document, dict):
            return cap_document
    return None


async def _state_file(
    container, *, flow_id: str, context_id: str, file_name: str, require_document: bool = False
) -> dict[str, Any]:
    session_id = f"{flow_id}:{context_id}"
    fallback: dict[str, Any] | None = None
    for _ in range(30):
        state = await container.workflow_runtime.get_state(session_id)
        if state:
            files_raw = (
                state.get("files") if isinstance(state, dict) else getattr(state, "files", [])
            )
            files = files_raw if isinstance(files_raw, list) else []
            matches = [
                file_item
                for item in files
                if (file_item := _file_item_dict(item)) is not None
                and file_item.get("original_name") == file_name
            ]
            for item in reversed(matches):
                if _file_document(item):
                    return item
            if matches:
                fallback = matches[-1]
                if not require_document:
                    return fallback
        await asyncio.sleep(0.1)
    if fallback is not None and (not require_document):
        return fallback
    raise AssertionError(f"File {file_name!r} was not persisted in state {session_id!r}")


async def _download_flow_file(flows_client_http, headers: dict[str, str], file_id: str) -> bytes:
    response = await flows_client_http.get(
        f"/frontend/api/v1/files/download/{file_id}", headers=headers
    )
    assert response.status_code == 200, response.text
    return response.content


@pytest.mark.asyncio
async def test_agent_edits_uploaded_docx_with_documents_tools(
    flows_client_http,
    office_client_http,
    office_service,
    mock_llm_redis,
    auth_headers_system,
    container,
    unique_id: str,
) -> None:
    headers = _headers(auth_headers_system)
    catalog_id = await _create_catalog(office_client_http, headers, unique_id)
    flow_id = f"documents_docx_{unique_id}"
    context_id = f"ctx-docx-{uuid.uuid4().hex}"
    file_name = f"agent-doc-{unique_id}.docx"
    await mock_llm_redis(
        [
            {
                "type": "tool_call",
                "tool": "documents_open_file",
                "args": {"file_name": file_name, "catalog_id": catalog_id},
            },
            {
                "type": "tool_call",
                "tool": "documents_replace_text",
                "args": {
                    "file_name": file_name,
                    "catalog_id": catalog_id,
                    "find": "ORIGINAL_TOKEN",
                    "replace": "UPDATED_TOKEN",
                    "match_case": True,
                },
            },
            {
                "type": "tool_call",
                "tool": "documents_append_text",
                "args": {
                    "file_name": file_name,
                    "catalog_id": catalog_id,
                    "text": "APPENDED_TOKEN",
                },
            },
            {"type": "text", "content": "Document updated."},
        ]
    )
    await _create_flow(flows_client_http, headers, flow_id)
    try:
        await _send_file_message(
            flows_client_http,
            headers,
            flow_id=flow_id,
            context_id=context_id,
            file_name=file_name,
            mime_type=DOCX_MIME,
            file_bytes=_docx_bytes("Before ORIGINAL_TOKEN", "Second paragraph"),
        )
        file_item = await _state_file(
            container,
            flow_id=flow_id,
            context_id=context_id,
            file_name=file_name,
            require_document=True,
        )
        document = _file_document(file_item)
        assert document
        assert document["file_id"] == file_item["file_id"]
        assert document["catalog_id"] == catalog_id
        assert document["editor_url"].startswith("/documents/embed/edit/")
        stored = await _download_flow_file(flows_client_http, headers, file_item["file_id"])
        text = _docx_text(stored)
        assert "UPDATED_TOKEN" in text
        assert "APPENDED_TOKEN" in text
        assert "ORIGINAL_TOKEN" not in text
    finally:
        await flows_client_http.delete(f"/flows/api/v1/flows/{flow_id}", headers=headers)


@pytest.mark.asyncio
async def test_agent_updates_uploaded_xlsx_with_documents_tool(
    flows_client_http,
    office_client_http,
    office_service,
    mock_llm_redis,
    auth_headers_system,
    container,
    unique_id: str,
) -> None:
    headers = _headers(auth_headers_system)
    catalog_id = await _create_catalog(office_client_http, headers, f"xlsx-{unique_id}")
    flow_id = f"documents_xlsx_{unique_id}"
    context_id = f"ctx-xlsx-{uuid.uuid4().hex}"
    file_name = f"agent-sheet-{unique_id}.xlsx"
    await mock_llm_redis(
        [
            {
                "type": "tool_call",
                "tool": "documents_open_file",
                "args": {"file_name": file_name, "catalog_id": catalog_id},
            },
            {
                "type": "tool_call",
                "tool": "documents_update_cells",
                "args": {
                    "file_name": file_name,
                    "catalog_id": catalog_id,
                    "sheet": "Sheet1",
                    "cells": {"A1": "UPDATED_CELL", "B2": 42},
                },
            },
            {"type": "text", "content": "Spreadsheet updated."},
        ]
    )
    await _create_flow(flows_client_http, headers, flow_id)
    try:
        await _send_file_message(
            flows_client_http,
            headers,
            flow_id=flow_id,
            context_id=context_id,
            file_name=file_name,
            mime_type=XLSX_MIME,
            file_bytes=_xlsx_bytes(),
        )
        file_item = await _state_file(
            container,
            flow_id=flow_id,
            context_id=context_id,
            file_name=file_name,
            require_document=True,
        )
        document = _file_document(file_item)
        assert document
        assert document["file_id"] == file_item["file_id"]
        assert document["catalog_id"] == catalog_id
        assert document["editor_url"].startswith("/documents/embed/edit/")
        stored = await _download_flow_file(flows_client_http, headers, file_item["file_id"])
        wb = load_workbook(BytesIO(stored))
        ws = wb["Sheet1"]
        assert ws["A1"].value == "UPDATED_CELL"
        assert ws["B2"].value == 42
    finally:
        await flows_client_http.delete(f"/flows/api/v1/flows/{flow_id}", headers=headers)
