"""FileWriter: markdown и прочие режимы в байты целевого файла."""

from __future__ import annotations

import base64
import binascii
import mimetypes
import re
from collections.abc import Callable
from io import BytesIO
from pathlib import Path
from typing import ClassVar

import markdown
from docx import Document
from docx.shared import Inches
from openpyxl.drawing.image import Image as XLImage
from openpyxl.workbook import Workbook
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Flowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.platypus import Image as RLImage

from core.files.checksum import compute_content_checksum_sha256
from core.files.create_spec import FileCreateSpec, FilePostCreate
from core.files.models import FileRecord
from core.files.service import FilesService
from core.files.types import ext_to_category, ext_to_mime
from core.files.writer.content_kind import (
    SourceContent,
    classify_content,
    normalize_str_content,
)
from core.files.writer.exceptions import FileWriteError
from core.files.writer.image_fetch import fetch_url_bytes
from core.files.writer.md_parse import flatten_markdown_segments, parse_gfm_table
from core.files.writer.models import ContentKind, ContentMode, FileWriteResult, WriteOptions
from core.files.writer.pdf_unicode_font import register_pdf_unicode_font


class FileWriter:
    """
    Сборка файла (build_bytes) и опционально загрузка в хранилище (create_file).
    """

    _process_files_service: ClassVar[FilesService | None] = None
    _process_default_spec: ClassVar[FileCreateSpec | None] = None

    @classmethod
    def configure_process_upload(
        cls,
        *,
        files_service: FilesService,
        default_spec: FileCreateSpec,
    ) -> None:
        cls._process_files_service = files_service
        cls._process_default_spec = default_spec

    def __init__(
        self,
        options: WriteOptions | None = None,
        *,
        files_service: FilesService | None = None,
        default_spec: FileCreateSpec | None = None,
    ) -> None:
        self._default_options: WriteOptions = options or WriteOptions()
        resolved_files_service: FilesService | None
        resolved_default_spec: FileCreateSpec | None
        if files_service is None and default_spec is None:
            resolved_files_service = FileWriter._process_files_service
            resolved_default_spec = FileWriter._process_default_spec
        elif files_service is None or default_spec is None:
            raise ValueError(
                "files_service and default_spec must be set together in constructor, "
                + "or both omitted for process defaults from configure_process_upload"
            )
        else:
            resolved_files_service = files_service
            resolved_default_spec = default_spec
        self._files_service: FilesService | None = resolved_files_service
        self._default_spec: FileCreateSpec | None = resolved_default_spec

    @classmethod
    def bind_for_upload(
        cls,
        *,
        files_service: FilesService,
        default_spec: FileCreateSpec,
        options: WriteOptions | None = None,
    ) -> FileWriter:
        return cls(
            options=options,
            files_service=files_service,
            default_spec=default_spec,
        )

    def _opts(self, override: WriteOptions | None) -> WriteOptions:
        return override if override is not None else self._default_options

    def _fetcher(self, opts: WriteOptions) -> Callable[[str], tuple[bytes, str | None]]:
        def inner(url: str) -> tuple[bytes, str | None]:
            return fetch_url_bytes(
                url,
                max_bytes=opts.max_image_bytes,
                timeout_seconds=opts.http_timeout_seconds,
            )

        return inner

    def build_bytes(
        self,
        content: SourceContent,
        original_name: str,
        *,
        content_mode: ContentMode = "auto",
        options: WriteOptions | None = None,
    ) -> FileWriteResult:
        opts = self._opts(options)
        name = (original_name or "").strip()
        if not name:
            raise FileWriteError("Параметр original_name обязателен (имя файла с расширением)")
        ext = Path(name).suffix.lower()
        if not ext:
            raise FileWriteError(
                f"В original_name нужно расширение целевого файла, получено: {original_name!r}"
            )

        mode = content_mode
        if mode == "auto":
            if isinstance(content, bytes):
                kind = ContentKind.RAW
            else:
                kind = classify_content(content)
            resolved = {
                ContentKind.MARKDOWN: "markdown",
                ContentKind.BASE64: "base64",
                ContentKind.RAW: "raw",
            }[kind]
        else:
            resolved = mode

        if resolved == "base64":
            return self._build_base64(content, ext)
        if resolved == "raw":
            return self._build_raw(content, ext, opts)
        if resolved == "markdown":
            text, _ = normalize_str_content(content, opts.text_encoding)
            return self._build_markdown_pipeline(text, ext, opts)
        raise FileWriteError(f"Неизвестный content_mode: {content_mode!r}")

    async def _persist_upload(
        self,
        built: FileWriteResult,
        *,
        original_name: str,
        public: bool,
        create_spec: FileCreateSpec | None,
    ) -> FileRecord:
        if self._files_service is None:
            raise FileWriteError(
                "write: files_service not configured — call FileWriter.configure_process_upload at startup"
            )
        spec = create_spec if create_spec is not None else self._default_spec
        if spec is None:
            raise FileWriteError("write: FileCreateSpec is required")
        post_create = spec.post_create
        if post_create is None:
            post_create = FilePostCreate(is_public=public)
        elif post_create.is_public != public:
            post_create = post_create.model_copy(update={"is_public": public})
            spec = spec.model_copy(update={"post_create": post_create})
        return await self._files_service.create(
            spec,
            built.data,
            original_name=original_name,
            content_type=built.content_type,
            content_sha256_hex=built.checksum_sha256_hex,
        )

    async def write(
        self,
        *,
        content: SourceContent,
        original_name: str,
        content_mode: ContentMode = "auto",
        public: bool = True,
        create_spec: FileCreateSpec | None = None,
        text_encoding: str = "utf-8",
        max_image_bytes: int = 15 * 1024 * 1024,
        http_timeout_seconds: float = 30.0,
        pdf_max_image_width_pt: float = 400.0,
        docx_image_width_inches: float = 5.0,
    ) -> FileRecord:
        opts = WriteOptions(
            text_encoding=text_encoding,
            max_image_bytes=max_image_bytes,
            http_timeout_seconds=http_timeout_seconds,
            pdf_max_image_width_pt=pdf_max_image_width_pt,
            docx_image_width_inches=docx_image_width_inches,
        )
        built = self.build_bytes(
            content, original_name, content_mode=content_mode, options=opts
        )
        return await self._persist_upload(
            built,
            original_name=original_name,
            public=public,
            create_spec=create_spec,
        )

    async def create_file(
        self,
        *,
        content: SourceContent,
        original_name: str,
        content_mode: ContentMode = "auto",
        options: WriteOptions | None = None,
        public: bool = True,
        create_spec: FileCreateSpec | None = None,
    ) -> FileRecord:
        if options is not None:
            built = self.build_bytes(
                content, original_name, content_mode=content_mode, options=options
            )
            return await self._persist_upload(
                built,
                original_name=original_name,
                public=public,
                create_spec=create_spec,
            )
        return await self.write(
            content=content,
            original_name=original_name,
            content_mode=content_mode,
            public=public,
            create_spec=create_spec,
        )

    def _build_base64(
        self,
        content: SourceContent,
        ext: str,
    ) -> FileWriteResult:
        if isinstance(content, bytes):
            raw_b = content
        else:
            compact = "".join(content.split())
            try:
                raw_b = base64.b64decode(compact, validate=True)
            except (ValueError, binascii.Error) as exc:
                raise FileWriteError("Невалидная base64-строка") from exc
        mime = _mime_for_extension(ext)
        return FileWriteResult(
            data=raw_b,
            content_type=mime,
            conversion_applied=False,
            checksum_sha256_hex=compute_content_checksum_sha256(raw_b),
        )

    def _build_raw(
        self,
        content: SourceContent,
        ext: str,
        opts: WriteOptions,
    ) -> FileWriteResult:
        if isinstance(content, bytes):
            raw_b = content
        else:
            raw_b = content.encode(opts.text_encoding)
        mime = _mime_for_extension(ext)
        return FileWriteResult(
            data=raw_b,
            content_type=mime,
            conversion_applied=False,
            checksum_sha256_hex=compute_content_checksum_sha256(raw_b),
        )

    def _build_markdown_pipeline(
        self,
        md_text: str,
        ext: str,
        opts: WriteOptions,
    ) -> FileWriteResult:
        if ext in (".md", ".txt"):
            data = md_text.encode(opts.text_encoding)
            mime = "text/markdown" if ext == ".md" else "text/plain; charset=utf-8"
            return FileWriteResult(
                data=data,
                content_type=mime,
                conversion_applied=True,
                checksum_sha256_hex=compute_content_checksum_sha256(data),
            )
        if ext == ".html":
            data = self._md_to_html_bytes(md_text, opts)
            return FileWriteResult(
                data=data,
                content_type="text/html; charset=utf-8",
                conversion_applied=True,
                checksum_sha256_hex=compute_content_checksum_sha256(data),
            )
        if ext == ".pdf":
            data = self._md_to_pdf_bytes(md_text, opts)
            return FileWriteResult(
                data=data,
                content_type="application/pdf",
                conversion_applied=True,
                checksum_sha256_hex=compute_content_checksum_sha256(data),
            )
        if ext == ".docx":
            data = self._md_to_docx_bytes(md_text, opts)
            return FileWriteResult(
                data=data,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                conversion_applied=True,
                checksum_sha256_hex=compute_content_checksum_sha256(data),
            )
        if ext == ".xlsx":
            data = self._md_to_xlsx_bytes(md_text, opts)
            return FileWriteResult(
                data=data,
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                conversion_applied=True,
                checksum_sha256_hex=compute_content_checksum_sha256(data),
            )
        raise FileWriteError(
            f"Конвертация markdown в формат {ext!r} не поддерживается. "
            + "Доступно: .html, .pdf, .docx, .xlsx, .md, .txt"
        )

    def _md_to_html_bytes(self, md_text: str, opts: WriteOptions) -> bytes:
        fetch = self._fetcher(opts)
        body = markdown.markdown(
            md_text,
            extensions=["tables", "fenced_code", "nl2br"],
        )
        embedded = _embed_http_images_in_html(body, fetch)
        full = (
            "<!DOCTYPE html>\n<html><head><meta charset=\"utf-8\"></head>"
            f"<body>\n{embedded}\n</body></html>"
        )
        return full.encode(opts.text_encoding)

    def _md_to_docx_bytes(self, md_text: str, opts: WriteOptions) -> bytes:
        fetch = self._fetcher(opts)
        doc = Document()
        w_in = opts.docx_image_width_inches
        for seg in flatten_markdown_segments(md_text):
            if seg["kind"] == "text":
                for block in seg["raw"].split("\n\n"):
                    line = block.strip()
                    if line:
                        _ = doc.add_paragraph(line)
            elif seg["kind"] == "table":
                grid = parse_gfm_table(seg["raw"])
                ncols = max(len(r) for r in grid)
                table = doc.add_table(rows=len(grid), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(grid):
                    for ci in range(ncols):
                        val = row[ci] if ci < len(row) else ""
                        table.rows[ri].cells[ci].text = val
            elif seg["kind"] == "image":
                data, _mime = fetch(seg["url"])
                _ = doc.add_picture(BytesIO(data), width=Inches(w_in))
        bio = BytesIO()
        doc.save(bio)
        return bio.getvalue()

    def _md_to_xlsx_bytes(self, md_text: str, opts: WriteOptions) -> bytes:
        fetch = self._fetcher(opts)
        wb = Workbook()
        ws = wb.active
        assert ws is not None
        row_idx = 1
        for seg in flatten_markdown_segments(md_text):
            if seg["kind"] == "text":
                for block in seg["raw"].split("\n\n"):
                    line = block.strip()
                    if line:
                        _ = ws.cell(row=row_idx, column=1, value=line)
                        row_idx += 1
            elif seg["kind"] == "table":
                grid = parse_gfm_table(seg["raw"])
                for row in grid:
                    for c_idx, val in enumerate(row, start=1):
                        _ = ws.cell(row=row_idx, column=c_idx, value=val)
                    row_idx += 1
                row_idx += 1
            elif seg["kind"] == "image":
                data, _mime = fetch(seg["url"])
                img = XLImage(BytesIO(data))
                _ = ws.add_image(img, f"A{row_idx}")
                row_idx += 20
        bio = BytesIO()
        wb.save(bio)
        return bio.getvalue()

    def _md_to_pdf_bytes(self, md_text: str, opts: WriteOptions) -> bytes:
        fetch = self._fetcher(opts)
        font_name = register_pdf_unicode_font()
        styles = getSampleStyleSheet()
        normal = ParagraphStyle(
            name="WriterPdfBody",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=12,
        )
        story: list[Flowable] = []
        max_w = opts.pdf_max_image_width_pt

        segments = flatten_markdown_segments(md_text)
        if not segments:
            story.append(Paragraph(" ", normal))

        for seg in segments:
            if seg["kind"] == "text":
                for block in seg["raw"].split("\n\n"):
                    line = block.strip()
                    if line:
                        story.append(Paragraph(_escape_reportlab_text(line), normal))
                        story.append(Spacer(1, 0.15 * inch))
            elif seg["kind"] == "table":
                raw_grid = _pad_table_rows(parse_gfm_table(seg["raw"]))
                grid = [
                    [Paragraph(_escape_reportlab_text(c), normal) for c in row]
                    for row in raw_grid
                ]
                t = Table(grid)
                t.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ]
                    )
                )
                story.append(t)
                story.append(Spacer(1, 0.2 * inch))
            elif seg["kind"] == "image":
                data, _mime = fetch(seg["url"])
                pil_im = PILImage.open(BytesIO(data))
                iw, ih = pil_im.size
                scale = min(max_w / float(iw), (A4[1] / 2) / float(ih), 1.0)
                rw = iw * scale
                rh = ih * scale
                story.append(RLImage(BytesIO(data), width=rw, height=rh))
                story.append(Spacer(1, 0.2 * inch))

        bio = BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=A4)
        doc.build(story)
        return bio.getvalue()


def _mime_for_extension(ext: str) -> str:
    mime, _ = mimetypes.guess_type(f"name{ext}")
    if mime:
        return mime
    if ext_to_category(ext) is not None:
        return ext_to_mime(ext)
    raise FileWriteError(f"Не удалось сопоставить MIME тип расширению {ext!r}")


def _escape_reportlab_text(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pad_table_rows(grid: list[list[str]]) -> list[list[str]]:
    if not grid:
        return grid
    n = max(len(r) for r in grid)
    return [r + [""] * (n - len(r)) for r in grid]


_IMG_SRC_RE = re.compile(
    r'(<img\b[^>]*\bsrc=)(["\'])(https?://[^"\']+)\2',
    re.IGNORECASE,
)


def _embed_http_images_in_html(
    html: str,
    fetch: Callable[[str], tuple[bytes, str | None]],
) -> str:
    def repl(m: re.Match[str]) -> str:
        prefix, quote, src = m.group(1), m.group(2), m.group(3)
        data, mime = fetch(src)
        if not mime:
            guessed, _ = mimetypes.guess_type(src)
            mime = guessed or "application/octet-stream"
        b64 = base64.b64encode(data).decode("ascii")
        return f'{prefix}{quote}data:{mime};base64,{b64}{quote}'

    return _IMG_SRC_RE.sub(repl, html)
