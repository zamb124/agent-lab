"""Deterministic byte-level mutations for Office-backed files."""

from __future__ import annotations

import csv
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell
from openpyxl.utils.cell import coordinate_to_tuple


class DocumentMutationError(ValueError):
    """Raised when a mutation is not supported or cannot be applied."""


def _ext(name: str) -> str:
    return Path(name or "").suffix.lower()


def _replace_text(value: str, find: str, replace: str, *, match_case: bool) -> tuple[str, int]:
    if find == "":
        raise DocumentMutationError("find must not be empty")
    if match_case:
        count = value.count(find)
        return value.replace(find, replace), count
    lower_value = value.lower()
    lower_find = find.lower()
    out: list[str] = []
    pos = 0
    count = 0
    step = len(find)
    while True:
        idx = lower_value.find(lower_find, pos)
        if idx < 0:
            out.append(value[pos:])
            break
        out.append(value[pos:idx])
        out.append(replace)
        pos = idx + step
        count += 1
    return "".join(out), count


def _mutate_docx_replace(data: bytes, find: str, replace: str, *, match_case: bool) -> tuple[bytes, int]:
    doc = Document(BytesIO(data))
    count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            next_text, changed = _replace_text(run.text, find, replace, match_case=match_case)
            if changed:
                run.text = next_text
                count += changed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        next_text, changed = _replace_text(run.text, find, replace, match_case=match_case)
                        if changed:
                            run.text = next_text
                            count += changed
    out = BytesIO()
    doc.save(out)
    return out.getvalue(), count


def _mutate_docx_append(data: bytes, text: str) -> bytes:
    doc = Document(BytesIO(data))
    for line in text.splitlines() or [""]:
        doc.add_paragraph(line)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _mutate_xlsx_replace(data: bytes, find: str, replace: str, *, match_case: bool) -> tuple[bytes, int]:
    wb = load_workbook(BytesIO(data))
    count = 0
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell, Cell):
                    continue
                if isinstance(cell.value, str):
                    next_text, changed = _replace_text(cell.value, find, replace, match_case=match_case)
                    if changed:
                        cell.value = next_text
                        count += changed
    out = BytesIO()
    wb.save(out)
    return out.getvalue(), count


def replace_text_in_document(
    *,
    data: bytes,
    original_name: str,
    find: str,
    replace: str,
    match_case: bool,
) -> tuple[bytes, int]:
    ext = _ext(original_name)
    if ext == ".docx":
        return _mutate_docx_replace(data, find, replace, match_case=match_case)
    if ext == ".xlsx":
        return _mutate_xlsx_replace(data, find, replace, match_case=match_case)
    if ext in {".txt", ".csv"}:
        text = data.decode("utf-8-sig")
        next_text, count = _replace_text(text, find, replace, match_case=match_case)
        return next_text.encode("utf-8"), count
    raise DocumentMutationError(f"replace_text is not supported for {ext or original_name!r}")


def append_text_to_document(*, data: bytes, original_name: str, text: str) -> bytes:
    ext = _ext(original_name)
    if ext == ".docx":
        return _mutate_docx_append(data, text)
    if ext in {".txt", ".csv"}:
        src = data.decode("utf-8-sig")
        separator = "" if src.endswith(("\n", "\r")) or not src else "\n"
        return f"{src}{separator}{text}".encode("utf-8")
    raise DocumentMutationError(f"append_text is not supported for {ext or original_name!r}")


def update_spreadsheet_cells(
    *,
    data: bytes,
    original_name: str,
    sheet: str | None,
    cells: dict[str, Any],
) -> bytes:
    if not cells:
        raise DocumentMutationError("cells must not be empty")
    ext = _ext(original_name)
    if ext == ".xlsx":
        wb = load_workbook(BytesIO(data))
        if sheet:
            if sheet not in wb.sheetnames:
                raise DocumentMutationError(f"sheet not found: {sheet}")
            ws = wb[sheet]
        else:
            ws = wb.active
            if ws is None:
                raise DocumentMutationError("active sheet not found")
        for coord, value in cells.items():
            ws[str(coord).upper()] = value
        out = BytesIO()
        wb.save(out)
        return out.getvalue()
    if ext == ".csv":
        text = data.decode("utf-8-sig")
        rows = list(csv.reader(StringIO(text)))
        for coord, value in cells.items():
            row_idx, col_idx = coordinate_to_tuple(str(coord).upper())
            while len(rows) < row_idx:
                rows.append([])
            row = rows[row_idx - 1]
            while len(row) < col_idx:
                row.append("")
            row[col_idx - 1] = "" if value is None else str(value)
        out = StringIO()
        writer = csv.writer(out)
        writer.writerows(rows)
        return out.getvalue().encode("utf-8")
    raise DocumentMutationError(f"update_cells is not supported for {ext or original_name!r}")
